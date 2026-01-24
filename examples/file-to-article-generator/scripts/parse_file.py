#!/usr/bin/env python3
"""
文件解析脚本：解析PDF/Word/图片文件，提取文本内容，提取并保存图片到本地
"""

import sys
import json
import os
import argparse
from pathlib import Path
from typing import Dict, List, Any
import base64


def parse_pdf(file_path: str, output_dir: str) -> Dict[str, Any]:
    """
    解析PDF文件，提取文本内容和图片信息，保存图片到本地
    
    Args:
        file_path: PDF文件路径
        output_dir: 输出目录
        
    Returns:
        包含文本和图片信息的字典
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {
            "error": "PyMuPDF未安装，无法解析PDF文件。请安装: pip install PyMuPDF==1.23.8"
        }
    
    result = {
        "file_type": "pdf",
        "text_content": "",
        "images": [],
        "metadata": {}
    }
    
    # 创建图片输出目录
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    try:
        doc = fitz.open(file_path)
        
        # 提取元数据
        metadata = doc.metadata
        if metadata:
            result["metadata"] = {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "page_count": len(doc)
            }
        else:
            result["metadata"]["page_count"] = len(doc)
        
        # 提取文本内容和图片
        text_content = []
        image_count = 0
        
        for page_num, page in enumerate(doc):
            # 提取文本
            text = page.get_text()
            if text:
                text_content.append(f"--- 第{page_num + 1}页 ---\n{text}\n")
            
            # 提取图片
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # 保存图片
                    image_filename = f"image_{image_count + 1:03d}.{image_ext}"
                    image_path = os.path.join(images_dir, image_filename)
                    
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    image_count += 1
                    
                    result["images"].append({
                        "index": image_count,
                        "description": f"PDF第{page_num + 1}页的第{img_index + 1}张图片",
                        "local_path": f"./images/{image_filename}",
                        "page": page_num + 1,
                        "format": image_ext
                    })
                except Exception as e:
                    # 图片提取失败，跳过
                    continue
        
        result["text_content"] = "\n".join(text_content)
        
        doc.close()
        
    except Exception as e:
        return {
            "error": f"解析PDF文件失败: {str(e)}"
        }
    
    return result


def parse_docx(file_path: str, output_dir: str) -> Dict[str, Any]:
    """
    解析Word文档，提取文本内容和图片信息，保存图片到本地
    
    Args:
        file_path: Word文件路径
        output_dir: 输出目录
        
    Returns:
        包含文本和图片信息的字典
    """
    try:
        from docx import Document
    except ImportError:
        return {
            "error": "python-docx未安装，无法解析Word文件。请安装: pip install python-docx==0.8.11"
        }
    
    result = {
        "file_type": "docx",
        "text_content": "",
        "images": [],
        "metadata": {}
    }
    
    # 创建图片输出目录
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    try:
        doc = Document(file_path)
        
        # 提取元数据
        core_props = doc.core_properties
        result["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "paragraph_count": len(doc.paragraphs)
        }
        
        # 提取文本内容和图片
        text_parts = []
        image_count = 0
        image_counter = 0
        
        # 遍历段落和表格提取图片
        for part in doc.part.iter_parts():
            if hasattr(part, "image"):
                try:
                    # 提取图片
                    image_data = part.blob
                    
                    # 保存图片
                    image_filename = f"image_{image_counter + 1:03d}.png"
                    image_path = os.path.join(images_dir, image_filename)
                    
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    
                    image_counter += 1
                    image_count += 1
                    
                    result["images"].append({
                        "index": image_count,
                        "description": f"Word文档中的第{image_count}张图片",
                        "local_path": f"./images/{image_filename}",
                        "format": "png"
                    })
                except Exception as e:
                    # 图片提取失败，跳过
                    continue
        
        # 遍历段落提取文本
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # 检查表格中的文本
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[表格{table_idx+1} 行{row_idx+1} 列{cell_idx+1}]: {para.text.strip()}")
        
        result["text_content"] = "\n".join(text_parts)
        
    except Exception as e:
        return {
            "error": f"解析Word文件失败: {str(e)}"
        }
    
    return result


def parse_image(file_path: str, output_dir: str) -> Dict[str, Any]:
    """
    解析图片文件，提取基本信息，保存图片到本地
    
    Args:
        file_path: 图片文件路径
        output_dir: 输出目录
        
    Returns:
        包含图片信息的字典
    """
    try:
        from PIL import Image
    except ImportError:
        return {
            "error": "Pillow未安装，无法解析图片文件。请安装: pip install Pillow==10.0.0"
        }
    
    result = {
        "file_type": "image",
        "text_content": "[图片文件，未进行OCR识别]",
        "images": [],
        "metadata": {}
    }
    
    # 创建图片输出目录
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    try:
        with Image.open(file_path) as img:
            # 获取图片格式
            file_ext = img.format.lower() if img.format else "png"
            
            # 保存图片到输出目录
            image_filename = f"image_001.{file_ext}"
            image_path = os.path.join(images_dir, image_filename)
            img.save(image_path)
            
            result["metadata"] = {
                "format": img.format,
                "size": img.size,
                "mode": img.mode,
                "description": f"{img.size[0]}x{img.size[1]} {img.format}图片"
            }
            
            result["images"].append({
                "index": 0,
                "description": result["metadata"]["description"],
                "local_path": f"./images/{image_filename}",
                "format": file_ext,
                "size": img.size
            })
    except Exception as e:
        return {
            "error": f"解析图片文件失败: {str(e)}"
        }
    
    return result


def main():
    # 支持两种输入方式：
    # 1. 命令行参数: python parse_file.py --file xxx.pdf --output ./output
    # 2. stdin JSON: echo '{"file_path": "xxx.pdf"}' | python parse_file.py

    file_path = None
    output_dir = "/home/gem/output"  # 沙箱默认输出目录

    # 尝试从 stdin 读取 JSON
    if not sys.stdin.isatty():
        try:
            stdin_data = sys.stdin.read().strip()
            if stdin_data:
                input_json = json.loads(stdin_data)
                file_path = input_json.get("file_path")
                output_dir = input_json.get("output_dir", output_dir)
        except json.JSONDecodeError:
            pass

    # 如果 stdin 没有提供，使用命令行参数
    if not file_path:
        parser = argparse.ArgumentParser(description='解析PDF/Word/图片文件，提取并保存图片')
        parser.add_argument('--file', '-f', required=True, help='要解析的文件路径')
        parser.add_argument('--output', '-o', default=output_dir, help='输出目录（默认：/home/gem/output）')

        args = parser.parse_args()
        file_path = args.file
        output_dir = args.output
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(json.dumps({
            "error": f"文件不存在: {file_path}"
        }, ensure_ascii=False, indent=2))
        sys.exit(1)
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 根据扩展名判断文件类型
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        result = parse_pdf(file_path, output_dir)
    elif file_ext in ['.docx', '.doc']:
        result = parse_docx(file_path, output_dir)
    elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
        result = parse_image(file_path, output_dir)
    else:
        result = {
            "error": f"不支持的文件格式: {file_ext}。支持的格式: PDF, DOCX, JPG, PNG, GIF, BMP, WEBP"
        }
    
    # 输出结果（JSON格式）
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
