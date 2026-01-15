#!/usr/bin/env python3
"""
Read docx file content.

Extracts text, tables, and metadata from Word documents.
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_tables(doc) -> list[dict]:
    """Extract tables from document."""
    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({
            "index": i,
            "rows": len(rows),
            "cols": len(rows[0]) if rows else 0,
            "data": rows
        })
    return tables


def main():
    """Main entry point."""
    # Read input
    stdin_content = sys.stdin.read().strip()

    try:
        input_data = json.loads(stdin_content)
        file_path = input_data.get("file_path", "")
    except json.JSONDecodeError:
        file_path = stdin_content

    if not file_path:
        print(json.dumps({"error": "No file path provided"}, ensure_ascii=False))
        return

    path = Path(file_path).expanduser()

    if not path.exists():
        print(json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False))
        return

    if not HAS_DOCX:
        # Fallback: just report file info
        print(json.dumps({
            "status": "warning",
            "message": "python-docx not installed. Install with: pip install python-docx",
            "file_path": str(path),
            "file_size": path.stat().st_size,
        }, ensure_ascii=False, indent=2))
        return

    try:
        doc = Document(path)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })

        # Extract tables
        tables = extract_tables(doc)

        # Get full text
        full_text = "\n\n".join(p["text"] for p in paragraphs)

        result = {
            "status": "success",
            "file_path": str(path),
            "file_name": path.name,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "content": full_text,
            "paragraphs": paragraphs[:20],  # Limit for output size
            "tables": tables[:5],  # Limit tables
        }

        print(json.dumps(result, ensure_ascii=False, indent=2))

    except Exception as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
