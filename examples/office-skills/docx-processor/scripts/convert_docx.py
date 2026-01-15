#!/usr/bin/env python3
"""
Convert docx to other formats.
"""

import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def to_markdown(doc) -> str:
    """Convert document to markdown."""
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"

        # Convert heading styles
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip())
                lines.append("#" * level + " " + text)
            except ValueError:
                lines.append("## " + text)
        elif style == "Title":
            lines.append("# " + text)
        elif style == "List Paragraph":
            lines.append("- " + text)
        else:
            lines.append(text)

        lines.append("")

    # Convert tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            # Add header separator after first row
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
            lines.extend(rows)
            lines.append("")

    return "\n".join(lines)


def main():
    """Main entry point."""
    stdin_content = sys.stdin.read().strip()

    try:
        input_data = json.loads(stdin_content)
        file_path = input_data.get("file_path", "")
        output_format = input_data.get("output_format", "markdown")
    except json.JSONDecodeError:
        # Simple path input
        file_path = stdin_content
        output_format = "markdown"

    if not file_path:
        print(json.dumps({"error": "No file path provided"}, ensure_ascii=False))
        return

    path = Path(file_path).expanduser()

    if not path.exists():
        print(json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False))
        return

    if not HAS_DOCX:
        print(json.dumps({
            "error": "python-docx not installed. Install with: pip install python-docx"
        }, ensure_ascii=False))
        return

    try:
        doc = Document(path)

        if output_format == "markdown":
            content = to_markdown(doc)
            ext = ".md"
        else:  # txt
            content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            ext = ".txt"

        # Save output
        output_dir = Path("/tmp/openskills-converted")
        output_dir.mkdir(parents=True, exist_ok=True)

        output_name = path.stem + ext
        output_path = output_dir / output_name
        output_path.write_text(content, encoding="utf-8")

        result = {
            "status": "success",
            "input_file": str(path),
            "output_file": str(output_path),
            "output_format": output_format,
            "content_preview": content[:500] + "..." if len(content) > 500 else content,
            "timestamp": datetime.now().isoformat(),
        }

        print(json.dumps(result, ensure_ascii=False, indent=2))

    except Exception as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
