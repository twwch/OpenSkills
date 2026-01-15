#!/usr/bin/env python3
"""
Office Skills Demo

演示如何使用 OpenSkills SDK 处理 Word 和 Excel 文件。
API Key 从环境变量获取。

环境变量:
    OPENAI_API_KEY: OpenAI API Key
    OPENAI_BASE_URL: (可选) API Base URL，默认 https://api.openai.com/v1
    OPENAI_MODEL: (可选) 模型名称，默认 gpt-4
"""

import asyncio
import os
import sys
from pathlib import Path

# 添加项目根目录到 Python 路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))


# ============================================================
# 示例数据：创建测试用的 docx 和 xlsx 文件
# ============================================================

def create_sample_docx(output_path: Path) -> bool:
    """创建示例 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # 添加标题
        doc.add_heading('Q1 季度销售报告', 0)

        # 添加段落
        doc.add_paragraph('本报告总结了2024年第一季度的销售情况。')

        doc.add_heading('销售概览', level=1)
        doc.add_paragraph('第一季度总销售额达到 1,250 万元，同比增长 15%。')
        doc.add_paragraph('主要增长来自华东和华南区域。')

        doc.add_heading('区域表现', level=1)

        # 添加表格
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '区域'
        header_cells[1].text = '销售额(万元)'
        header_cells[2].text = '同比增长'

        # 数据
        data = [
            ('华东', '450', '+18%'),
            ('华南', '380', '+22%'),
            ('华北', '420', '+8%'),
        ]

        for i, (region, sales, growth) in enumerate(data, 1):
            row = table.rows[i].cells
            row[0].text = region
            row[1].text = sales
            row[2].text = growth

        doc.add_heading('总结', level=1)
        doc.add_paragraph('整体表现良好，建议继续加强华东和华南市场的投入。')

        doc.save(output_path)
        return True
    except ImportError:
        print("[警告] python-docx 未安装，跳过创建示例 docx")
        return False


def create_sample_xlsx(output_path: Path) -> bool:
    """创建示例 Excel 文件"""
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "销售数据"

        # 表头
        headers = ['月份', '华东', '华南', '华北', '总计']
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)

        # 数据
        data = [
            ['1月', 150, 120, 140, 410],
            ['2月', 145, 125, 135, 405],
            ['3月', 155, 135, 145, 435],
        ]

        for row_idx, row_data in enumerate(data, 2):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # 添加第二个工作表
        ws2 = wb.create_sheet("产品分类")
        ws2.append(['产品', '销量', '单价', '收入'])
        ws2.append(['产品A', 1000, 50, 50000])
        ws2.append(['产品B', 800, 80, 64000])
        ws2.append(['产品C', 1200, 30, 36000])

        wb.save(output_path)
        return True
    except ImportError:
        print("[警告] openpyxl 未安装，跳过创建示例 xlsx")
        return False


# ============================================================
# Demo 函数
# ============================================================

async def demo_docx_processing():
    """演示 Word 文档处理"""
    from openskills import create_agent

    print("\n" + "=" * 60)
    print("Demo 1: Word 文档处理")
    print("=" * 60)

    # 从环境变量获取配置
    api_key = os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    model = os.environ.get("OPENAI_MODEL", "gpt-4")

    if not api_key:
        print("[错误] 请设置 OPENAI_API_KEY 环境变量")
        print("  export OPENAI_API_KEY=your-api-key")
        return

    # 创建示例文件
    sample_dir = Path("/tmp/openskills-samples")
    sample_dir.mkdir(parents=True, exist_ok=True)
    docx_path = sample_dir / "sales_report.docx"

    if not create_sample_docx(docx_path):
        print("[跳过] 无法创建示例文件")
        return

    print(f"\n[示例文件] {docx_path}")

    # 回调函数
    def on_reference_loaded(ref_path: str, content: str):
        print(f"\n[Reference 加载] {ref_path}")

    def on_script_executed(script_name: str, result: str):
        print(f"\n[脚本执行] {script_name}")
        # 解析并美化输出
        try:
            import json
            data = json.loads(result)
            if data.get("status") == "success":
                print(f"  状态: 成功")
                if "content" in data:
                    print(f"  内容预览: {data['content'][:200]}...")
            else:
                print(f"  结果: {result[:200]}")
        except:
            print(f"  结果: {result[:200]}")

    # 创建 agent
    skills_path = Path(__file__).parent
    agent = await create_agent(
        skill_paths=[skills_path],
        api_key=api_key,
        base_url=base_url,
        model=model,
        auto_execute_scripts=True,
        on_reference_loaded=on_reference_loaded,
        on_script_executed=on_script_executed,
    )

    print(f"\n[发现 Skills] {agent.available_skills}")

    # 用户请求
    user_message = f"""请帮我读取并分析这个 Word 文档: {docx_path}

总结文档的主要内容，提取关键数据。"""

    print(f"\n[用户请求]\n{user_message}")
    print("\n[处理中...]")

    response = await agent.chat(user_message)

    print(f"\n[AI 回复]\n{response.content}")
    print(f"\n[使用的 Skill] {response.skill_used}")
    if response.references_loaded:
        print(f"[加载的 References] {response.references_loaded}")
    if response.scripts_executed:
        print(f"[执行的脚本] {response.scripts_executed}")


async def demo_excel_processing():
    """演示 Excel 文件处理"""
    from openskills import create_agent

    print("\n" + "=" * 60)
    print("Demo 2: Excel 文件处理")
    print("=" * 60)

    # 从环境变量获取配置
    api_key = os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    model = os.environ.get("OPENAI_MODEL", "gpt-4")

    if not api_key:
        print("[错误] 请设置 OPENAI_API_KEY 环境变量")
        return

    # 创建示例文件
    sample_dir = Path("/tmp/openskills-samples")
    sample_dir.mkdir(parents=True, exist_ok=True)
    xlsx_path = sample_dir / "sales_data.xlsx"

    if not create_sample_xlsx(xlsx_path):
        print("[跳过] 无法创建示例文件")
        return

    print(f"\n[示例文件] {xlsx_path}")

    # 回调函数
    def on_reference_loaded(ref_path: str, content: str):
        print(f"\n[Reference 加载] {ref_path}")

    def on_script_executed(script_name: str, result: str):
        print(f"\n[脚本执行] {script_name}")
        try:
            import json
            data = json.loads(result)
            if "sheets" in data:
                for sheet in data["sheets"]:
                    print(f"  工作表: {sheet['name']} ({sheet['rows']}行 x {sheet['cols']}列)")
            elif "summary" in data:
                print(f"  总计: {data['summary']['total_rows']} 行")
        except:
            print(f"  结果: {result[:200]}")

    # 创建 agent
    skills_path = Path(__file__).parent
    agent = await create_agent(
        skill_paths=[skills_path],
        api_key=api_key,
        base_url=base_url,
        model=model,
        auto_execute_scripts=True,
        on_reference_loaded=on_reference_loaded,
        on_script_executed=on_script_executed,
    )

    # 用户请求
    user_message = f"""请帮我分析这个 Excel 文件: {xlsx_path}

我想了解：
1. 文件包含哪些工作表
2. 各区域的销售数据统计
3. 哪个月份表现最好"""

    print(f"\n[用户请求]\n{user_message}")
    print("\n[处理中...]")

    response = await agent.chat(user_message)

    print(f"\n[AI 回复]\n{response.content}")
    print(f"\n[使用的 Skill] {response.skill_used}")
    if response.references_loaded:
        print(f"[加载的 References] {response.references_loaded}")
    if response.scripts_executed:
        print(f"[执行的脚本] {response.scripts_executed}")


async def demo_combined():
    """组合演示：同时处理 Word 和 Excel"""
    from openskills import create_agent

    print("\n" + "=" * 60)
    print("Demo 3: 组合处理 (Word + Excel)")
    print("=" * 60)

    api_key = os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    model = os.environ.get("OPENAI_MODEL", "gpt-4")

    if not api_key:
        print("[错误] 请设置 OPENAI_API_KEY 环境变量")
        return

    sample_dir = Path("/tmp/openskills-samples")
    docx_path = sample_dir / "sales_report.docx"
    xlsx_path = sample_dir / "sales_data.xlsx"

    # 确保文件存在
    create_sample_docx(docx_path)
    create_sample_xlsx(xlsx_path)

    def on_script_executed(script_name: str, result: str):
        print(f"  [执行] {script_name}")

    skills_path = Path(__file__).parent
    agent = await create_agent(
        skill_paths=[skills_path],
        api_key=api_key,
        base_url=base_url,
        model=model,
        auto_execute_scripts=True,
        on_script_executed=on_script_executed,
    )

    print(f"\n[可用 Skills] {agent.available_skills}")

    # 第一轮对话：处理 Excel
    print("\n--- 第一轮：分析 Excel ---")
    response1 = await agent.chat(f"读取 Excel 文件 {xlsx_path}，告诉我总销售额")
    print(f"回复: {response1.content[:300]}...")

    # 第二轮对话：处理 Word（需要重置或使用新 agent）
    agent.reset()
    print("\n--- 第二轮：分析 Word ---")
    response2 = await agent.chat(f"读取 Word 文档 {docx_path}，提取关键结论")
    print(f"回复: {response2.content[:300]}...")


async def main():
    """主函数"""
    print("=" * 60)
    print("OpenSkills Office Demo")
    print("=" * 60)
    print(f"\n环境变量配置:")
    print(f"  OPENAI_API_KEY: {'已设置' if os.environ.get('OPENAI_API_KEY') else '未设置'}")
    print(f"  OPENAI_BASE_URL: {os.environ.get('OPENAI_BASE_URL', '(默认)')}")
    print(f"  OPENAI_MODEL: {os.environ.get('OPENAI_MODEL', '(默认 gpt-4)')}")

    # 检查依赖
    print(f"\n依赖检查:")
    try:
        import docx
        print(f"  python-docx: 已安装")
    except ImportError:
        print(f"  python-docx: 未安装 (pip install python-docx)")

    try:
        import openpyxl
        print(f"  openpyxl: 已安装")
    except ImportError:
        print(f"  openpyxl: 未安装 (pip install openpyxl)")

    if not os.environ.get("OPENAI_API_KEY"):
        print("\n[提示] 请设置环境变量后运行:")
        print("  export OPENAI_API_KEY=your-api-key")
        print("  export OPENAI_BASE_URL=https://api.openai.com/v1  # 可选")
        print("  export OPENAI_MODEL=gpt-4  # 可选")
        return

    # 运行 demo
    # await demo_docx_processing()
    # await demo_excel_processing()
    await demo_combined()  # 取消注释运行组合演示


if __name__ == "__main__":
    asyncio.run(main())
